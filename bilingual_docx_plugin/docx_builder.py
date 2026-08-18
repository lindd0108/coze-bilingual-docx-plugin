from __future__ import annotations

import re
import uuid
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def safe_file_stem(value: str) -> str:
    stem = re.sub(r'[<>:"/\\|?*]', "_", value)
    stem = re.sub(r"[\x00-\x1f]", "_", stem).strip(" ._")
    return stem[:80] or "bilingual_docx"


def add_paragraph_text(cell, text: str, font_name: str, font_size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text or "").strip())
    run.font.name = font_name
    run.font.size = Pt(font_size)


def set_cell_text(cell, text: str, font_name: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def normalize_rows(rows: Any) -> List[Dict[str, str]]:
    if not isinstance(rows, list):
        return []
    normalized = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        source = str(row.get("source", "")).strip()
        target = str(row.get("target", "")).strip()
        if source or target:
            normalized.append({"source": source, "target": target})
    return normalized


def normalize_decision_logs(logs: Any) -> List[Dict[str, str]]:
    if not isinstance(logs, list):
        return []
    normalized = []
    for item in logs:
        if not isinstance(item, dict):
            continue
        log = {
            "item": str(item.get("item", "")).strip(),
            "decision": str(item.get("decision", "")).strip(),
            "reason": str(item.get("reason", "")).strip(),
            "memory": str(item.get("memory", "")).strip(),
            "review": str(item.get("review", "")).strip(),
        }
        if any(log.values()):
            normalized.append(log)
    return normalized


def build_docx(payload: Dict[str, Any], output_dir: Path) -> Dict[str, Any]:
    title = str(payload.get("title", "")).strip()
    rows = normalize_rows(payload.get("rows"))
    terms = payload.get("terms") if isinstance(payload.get("terms"), list) else []
    risks = payload.get("risks") if isinstance(payload.get("risks"), list) else []
    decision_logs = normalize_decision_logs(payload.get("decision_logs"))

    if not title:
        raise ValueError("title is required")
    if not rows:
        raise ValueError("rows must contain at least one non-empty row")

    output_dir.mkdir(exist_ok=True)
    file_name = f"{safe_file_stem(title)}_{uuid.uuid4().hex[:8]}.docx"
    output_path = output_dir / file_name

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0].cells
    set_cell_text(header[0], "中文原文", "SimSun", bold=True)
    set_cell_text(header[1], "英文译文", "Times New Roman", bold=True)

    for row in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        add_paragraph_text(cells[0], row["source"], "SimSun")
        add_paragraph_text(cells[1], row["target"], "Times New Roman")

    if terms:
        document.add_heading("术语说明", level=2)
        term_table = document.add_table(rows=1, cols=3)
        term_table.style = "Table Grid"
        for cell, text in zip(term_table.rows[0].cells, ["中文项", "推荐译法", "说明"]):
            set_cell_text(cell, text, "SimSun", bold=True)
        for term in terms:
            if not isinstance(term, dict):
                continue
            cells = term_table.add_row().cells
            add_paragraph_text(cells[0], str(term.get("term", "")), "SimSun")
            add_paragraph_text(cells[1], str(term.get("translation", "")), "Times New Roman")
            add_paragraph_text(cells[2], str(term.get("note", "")), "SimSun")

    if risks:
        document.add_heading("风险清单", level=2)
        for risk in risks:
            risk_text = str(risk).strip()
            if risk_text:
                document.add_paragraph(risk_text)

    if decision_logs:
        document.add_heading("译者决策日志", level=2)
        log_table = document.add_table(rows=1, cols=5)
        log_table.style = "Table Grid"
        headers = ["项目", "决策", "理由", "是否写入项目记忆", "是否需复核"]
        for cell, text in zip(log_table.rows[0].cells, headers):
            set_cell_text(cell, text, "SimSun", bold=True)
        for log in decision_logs:
            cells = log_table.add_row().cells
            add_paragraph_text(cells[0], log["item"], "SimSun", font_size=9)
            add_paragraph_text(cells[1], log["decision"], "SimSun", font_size=9)
            add_paragraph_text(cells[2], log["reason"], "SimSun", font_size=9)
            add_paragraph_text(cells[3], log["memory"], "SimSun", font_size=9)
            add_paragraph_text(cells[4], log["review"], "SimSun", font_size=9)

    document.save(output_path)
    return {
        "file_name": file_name,
        "output_path": output_path,
        "row_count": len(rows),
    }
